import os
import json
import zipfile
import shutil
from pathlib import Path
from typing import Dict, List
from datetime import date

from docx import Document
from docx.shared import Inches

from database_manager import get_articles_for_delivery

'''
Модуль финальной сборки.
Собирает сгенерированные за день статьи и изображения,
упаковывает их в DOCX-файлы и создает ZIP-архивы для каждого пользователя.
'''

# --- Конфигурация ---
OUTPUT_ZIP_DIR = "daily_zips"


# --- Вспомогательные функции ---

def create_docx(article_data: Dict, filepath: Path):
    """Создает и сохраняет DOCX-файл для одной статьи."""
    doc = Document()

    # Форматируем токены из JSON-строки в обычную строку
    try:
        tokens_list = json.loads(article_data['matched_tokens'])
        tokens_str = ", ".join(tokens_list).upper()
    except (json.JSONDecodeError, TypeError):
        tokens_str = "N/A"

    # Заполняем документ по шаблону
    doc.add_paragraph(f"TITLE: {article_data['title']}")
    doc.add_paragraph(f"CATEGORY: {article_data['category']}")
    doc.add_paragraph(f"TOKEN: {tokens_str}")
    doc.add_paragraph("ARTICLE:")
    doc.add_paragraph(article_data['content'])

    doc.save(filepath)


def sanitize_filename(name: str) -> str:
    """Очищает строку, чтобы она была безопасным именем файла."""
    return "".join(c for c in name if c.isalnum() or c in (' ', '_')).rstrip()


# --- Основная логика ---

def run_doc_zipper() -> Dict[int, str]:
    """
    Основная функция. Собирает статьи, создает DOCX и ZIP.
    Возвращает словарь {user_id: path_to_zip}.
    """
    print("  -> Запуск doc_zipper.py...")

    articles_to_deliver = get_articles_for_delivery()
    if not articles_to_deliver:
        print("     [INFO] Нет статей для сборки и доставки.")
        return {}

    # Создаем корневую папку для ZIP-архивов
    Path(OUTPUT_ZIP_DIR).mkdir(exist_ok=True)

    generated_zips = {}

    # Цикл по каждому пользователю, у которого есть статьи
    for user_id, articles in articles_to_deliver.items():
        username = articles[0].get('username', f'user_{user_id}')
        print(f"     Начинаю сборку для пользователя: {username} ({len(articles)} статей)")

        # Создаем временную папку для файлов этого пользователя
        temp_user_dir = Path(OUTPUT_ZIP_DIR) / f"temp_{user_id}"
        temp_user_dir.mkdir(exist_ok=True)

        try:
            # Цикл по каждой статье пользователя
            for article in articles:
                # 1. Создаем DOCX
                safe_title = sanitize_filename(article['title'])[:50]  # Ограничиваем длину
                docx_filename = f"{safe_title}.docx"
                docx_filepath = temp_user_dir / docx_filename
                create_docx(article, docx_filepath)

                # 2. Копируем изображение
                image_path_str = article.get('image_path')
                if image_path_str and Path(image_path_str).exists():
                    image_filename = f"{safe_title}.png"
                    image_filepath = temp_user_dir / image_filename
                    shutil.copy(image_path_str, image_filepath)
                else:
                    print(f"     [WARNING] Изображение для статьи '{safe_title}' не найдено.")

            # 3. Создаем ZIP-архив
            today_str = date.today().strftime('%Y-%m-%d')
            # TODO: Получить название стратегии из БД для более полного имени
            zip_filename = f"{username}_digest_{today_str}.zip"
            zip_filepath = Path(OUTPUT_ZIP_DIR) / zip_filename

            with zipfile.ZipFile(zip_filepath, 'w', zipfile.ZIP_DEFLATED) as zf:
                for file_to_zip in temp_user_dir.glob('*'):
                    zf.write(file_to_zip, arcname=file_to_zip.name)

            generated_zips[user_id] = str(zip_filepath)
            print(f"     [SUCCESS] ZIP-архив создан: {zip_filepath}")

        finally:
            # 4. Удаляем временную папку
            if temp_user_dir.exists():
                shutil.rmtree(temp_user_dir)

    return generated_zips


if __name__ == '__main__':
    print("--- Тестовый запуск Doc Zipper ---")
    final_zips = run_doc_zipper()
    if final_zips:
        print("\n--- Сборка успешно завершена. Созданы следующие архивы: ---")
        for user, path in final_zips.items():
            print(f"  - Пользователь {user}: {path}")
    else:
        print("\n--- Работа модуля Doc Zipper завершена (нечего было собирать). ---")